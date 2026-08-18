import streamlit as st
import requests
import pandas as pd
from collections import Counter
import re
import time
import io
from bs4 import BeautifulSoup
import docx
import PyPDF2
from gtts import gTTS
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from textblob import TextBlob
import base64

# Set up giao diện Streamlit
st.set_page_config(page_title="Radar SEO & Content Pro", page_icon="🚀", layout="wide")
st.title("🚀 Radar SEO & Content Pro (Bản All-in-One + Ultimate)")
st.markdown("Hệ sinh thái phân tích từ khóa, tình báo đối thủ và tự động hóa sản xuất nội dung bằng Trí Tuệ Nhân Tạo.")

# ================= KHỞI TẠO BỘ NHỚ TẠM =================
if 'history' not in st.session_state: st.session_state.history = []
if 'current_kw' not in st.session_state: st.session_state.current_kw = None
if 'current_text' not in st.session_state: st.session_state.current_text = ""
if 'current_file' not in st.session_state: st.session_state.current_file = ""
if 'ai_result' not in st.session_state: st.session_state.ai_result = ""
if 'sc_memory' not in st.session_state: st.session_state.sc_memory = ""
if 'sc_part' not in st.session_state: st.session_state.sc_part = 1
if 'yt_xray_data' not in st.session_state: st.session_state.yt_xray_data = None
if 'chat_history' not in st.session_state: st.session_state.chat_history = []
# Bộ nhớ cho Kho Prompt tùy chỉnh
if 'custom_prompts' not in st.session_state: st.session_state.custom_prompts = {}

# --- CẤU HÌNH API KEY TỰ ĐỘNG ---
default_gemini = st.secrets.get("GEMINI_API_KEY", "") if "GEMINI_API_KEY" in st.secrets else ""
default_yt = st.secrets.get("YOUTUBE_API_KEY", "") if "YOUTUBE_API_KEY" in st.secrets else ""

st.sidebar.header("🔑 Cấu Hình API")
gemini_api_key = st.sidebar.text_input("Gemini API Key:", value=default_gemini, type="password")
yt_api_key = st.sidebar.text_input("YouTube API v3 Key:", value=default_yt, type="password")

st.sidebar.markdown("---")
st.sidebar.header("🤖 Tùy chỉnh Model AI")

if st.sidebar.button("🔍 Quét Model Khả Dụng"):
    if not gemini_api_key: st.sidebar.error("Vui lòng nhập API Key trước!")
    else:
        with st.sidebar.status("Đang quét máy chủ Google..."):
            try:
                url_check = f"https://generativelanguage.googleapis.com/v1beta/models?key={gemini_api_key}"
                res_check = requests.get(url_check, timeout=10)
                if res_check.status_code == 200:
                    models = [m['name'].replace('models/', '') for m in res_check.json().get('models', []) if 'generateContent' in m.get('supportedGenerationMethods', [])]
                    st.sidebar.success("✅ Danh sách Model của bạn:")
                    for m in models: st.sidebar.code(m)
                else: st.sidebar.error(f"Lỗi: {res_check.status_code} - {res_check.text}")
            except Exception as e: st.sidebar.error(f"Lỗi kết nối: {e}")

ai_model_choice = st.sidebar.text_input("Nhập chính xác tên Model:", value="gemini-3.5-flash")

# TÍNH NĂNG 4: KHO PROMPT TÙY CHỈNH
st.sidebar.markdown("---")
st.sidebar.header("🗂️ Kho Prompt Bí Mật")
st.sidebar.caption("Tạo câu lệnh riêng. Dùng chữ **{tu_khoa}** để hệ thống tự chèn từ khóa vào.")
new_p_name = st.sidebar.text_input("Tên Hành động (Ví dụ: Viết kịch bản Facebook):")
new_p_content = st.sidebar.text_area("Nội dung Prompt:")
if st.sidebar.button("💾 Lưu Prompt"):
    if new_p_name and new_p_content:
        st.session_state.custom_prompts[new_p_name] = new_p_content
        st.sidebar.success("Đã lưu vào bộ nhớ!")

st.sidebar.markdown("---")
st.sidebar.header("⚙️ Cấu Hình Phân Tích")

# --- XỬ LÝ STOPWORDS ---
STOP_WORDS = {'the', 'and', 'to', 'of', 'a', 'in', 'for', 'is', 'on', 'that', 'by', 'this', 'with', 'i', 'you', 'it', 'not', 'or', 'be', 'are', 'from', 'at', 'as', 'your', 'how', 'what', 'why', 'do', 'can', 'my', 'we', 'about', 'an', 'if', 'will', 'https', 'http', 'com', 'www', 'youtube', 'watch', 'channel', 'video', 'v', 'org', 'net', 'và', 'là', 'của', 'cho', 'trong', 'với', 'các', 'những', 'đó', 'này', 'được', 'khi', 'về', 'có', 'không', 'như', 'đã', 'đang', 'sẽ', 'để', 'một', 'mọi', 'ra', 'vào', 'hoặc', 'vì', 'theo', 'tại', 'từ', 'nên', 'cần', 'nhưng', 'bị'}
uploaded_stopwords = st.sidebar.file_uploader("Tải lên file Stopwords (.txt)", type=['txt'])
if uploaded_stopwords:
    custom_words = uploaded_stopwords.read().decode('utf-8').splitlines()
    STOP_WORDS.update([w.strip().lower() for w in custom_words if w.strip()])

ngram_choice = st.sidebar.radio("Độ dài từ khóa:", ("Từ đơn", "Cụm 2 từ", "Cụm 3 từ"))
ngram_map = {"Từ đơn": 1, "Cụm 2 từ": 2, "Cụm 3 từ": 3}
selected_ngram = ngram_map[ngram_choice]

# --- CÁC HÀM XỬ LÝ CỐT LÕI ---
@st.cache_data(show_spinner=False)
def lam_sach_text(text):
    text = re.sub(r'https?://\S+|www\.\S+', '', text)
    text = re.sub(r'<[^>]+>', ' ', text)
    words = re.findall(r'\b[a-zA-Z0-9àáảãạăắằẳẵặâấầẩẫậèéẻẽẹêếềểễệđìíỉĩịòóỏõọôốồổỗộơớờởỡợùúủũụưứừửữựỳýỷỹỵ]+\b', text.lower())
    return words

def trich_xuat_tu_khoa(danh_sach_tu, loai_ngram=1, top_n=50):
    tu_hop_le = [w for w in danh_sach_tu if len(w) > 1 and w not in STOP_WORDS and not w.isdigit()]
    if loai_ngram == 1: tokens = tu_hop_le
    elif loai_ngram == 2: tokens = [' '.join(tu_hop_le[i:i+2]) for i in range(len(tu_hop_le)-1)]
    else: tokens = [' '.join(tu_hop_le[i:i+3]) for i in range(len(tu_hop_le)-2)]
    return Counter(tokens).most_common(top_n)

def phan_loai_y_dinh(kw):
    kw_lower = kw.lower()
    if any(w in kw_lower for w in ['mua', 'giá', 'bán', 'rẻ', 'bao', 'shop', 'lazada', 'shopee']): return "🛒 Giao dịch"
    if any(w in kw_lower for w in ['cách', 'hướng dẫn', 'là gì', 'tại sao', 'review']): return "📖 Thông tin"
    if any(w in kw_lower for w in ['facebook', 'youtube', 'tiki', 'login', 'web']): return "🧭 Điều hướng"
    return "💡 Khám phá"

def gom_nhom(kw):
    words = kw.split()
    if len(words) > 1: return words[0].capitalize()
    return "Từ Đơn"

def luu_ket_qua_vao_bo_nho(res_kw, full_text, ten_file):
    st.session_state.current_kw = res_kw
    st.session_state.current_text = full_text
    st.session_state.current_file = ten_file
    st.session_state.ai_result = ""
    st.session_state.sc_memory = ""
    st.session_state.sc_part = 1
    st.session_state.yt_xray_data = None
    st.session_state.chat_history = []

def ai_cham_diem_thumbnail(img_url, title):
    if not gemini_api_key: return "⚠️ Thiếu Gemini API Key."
    try:
        img_resp = requests.get(img_url, timeout=10)
        if img_resp.status_code != 200: return "⚠️ Lỗi tải ảnh."
        img_b64 = base64.b64encode(img_resp.content).decode('utf-8')
        model_path = ai_model_choice.strip().replace("models/", "")
        url_gen = f"https://generativelanguage.googleapis.com/v1beta/models/{model_path}:generateContent"
        headers = {'x-goog-api-key': gemini_api_key, 'Content-Type': 'application/json'}
        prompt = f"""Đánh giá Thumbnail và Tiêu đề: "{title}". Chấm điểm (/10) về màu sắc, text và CTR."""
        payload = {"contents": [{"parts": [{"text": prompt}, {"inline_data": {"mime_type": "image/jpeg", "data": img_b64}}]}]}
        res = requests.post(url_gen, headers=headers, json=payload, timeout=180)
        if res.status_code == 200: return res.json()['candidates'][0]['content']['parts'][0]['text']
        else: return f"⚠️ Lỗi Google ({res.status_code})."
    except Exception as e: return f"⚠️ Lỗi mạng: {e}"

# --- HÀM TẠO NỘI DUNG AI ĐA NĂNG ---
def xu_ly_ai_da_nang(che_do, tu_khoa_list, text_goc, is_continue=False, use_web=False):
    if not gemini_api_key:
        st.error("❌ Vui lòng dán Gemini API Key ở thanh bên trái!")
        return

    model_path = ai_model_choice.strip().replace("models/", "")
    url_gen = f"https://generativelanguage.googleapis.com/v1beta/models/{model_path}:generateContent"
    headers = {'x-goog-api-key': gemini_api_key, 'Content-Type': 'application/json'}

    if che_do == "🎬 Kịch Bản Phim Tài Liệu (Chuẩn Silent Capital)" and not is_continue:
        st.session_state.sc_memory = ""
        st.session_state.sc_part = 1

    # Tích hợp Custom Prompts
    if che_do in st.session_state.custom_prompts:
        prompt = st.session_state.custom_prompts[che_do].replace("{tu_khoa}", str(tu_khoa_list))
    elif che_do == "📝 Lập Dàn Ý SEO (Outline)": prompt = f"Dựa vào Top từ khóa: {tu_khoa_list}. Lập Dàn ý chi tiết chuẩn SEO."
    elif che_do == "✍️ Viết Bài Full Chuẩn SEO (1000+ từ)": prompt = f"Viết bài Blog chuẩn SEO dài khoảng 1000 từ dựa vào bộ từ khóa: {tu_khoa_list}."
    elif che_do == "🎬 Kịch Bản Video Ngắn (TikTok/Reels/Shorts)": prompt = f"Viết kịch bản Video dọc (dưới 60s) từ từ khóa: {tu_khoa_list}. Cấu trúc: Hook, Body, CTA."
    elif che_do == "🛒 Tối Ưu Gian Hàng TMĐT (Tiêu đề & Mô tả)": prompt = f"Viết 3 Tiêu đề giật tít và Mô tả sản phẩm đánh vào nỗi đau khách hàng từ: {tu_khoa_list}."
    elif che_do == "🕵️ Tái tạo Kịch Bản YouTube Đối Thủ (Phân tích Link)": prompt = f"Dữ liệu từ YouTube đối thủ: {text_goc[:5000]}. Viết Kịch Bản Video YouTube hoàn chỉnh, MỚI MẺ HƠN."
    elif che_do == "🧠 Phản Biện Kịch Bản (Script Doctor)": prompt = f"Phản biện khắt khe nội dung: {st.session_state.sc_memory if st.session_state.sc_memory else text_goc[:3000]}. Chỉ ra điểm mù logic."
    elif che_do == "📈 Dự Báo Thuật Toán YouTube (Pre-Publish)": prompt = f"Chấm điểm kịch bản dự kiến: {st.session_state.sc_memory if st.session_state.sc_memory else text_goc[:4000]}. Dự báo CTR, Retention."
    elif che_do == "🎬 Kịch Bản Phim Tài Liệu (Chuẩn Silent Capital)":
        prompt = f"""Biên kịch Phim tài liệu. Viết DUY NHẤT PART {st.session_state.sc_part} nối tiếp Kịch bản: {st.session_state.sc_memory}."""

    payload = {"contents": [{"role": "user", "parts": [{"text": prompt}]}]}
    
    # TÍNH NĂNG 2: MẮT THẦN INTERNET (GOOGLE SEARCH RETRIEVAL)
    if use_web:
        payload["tools"] = [{"googleSearch": {}}]

    with st.spinner(f"🤖 Đang gọi Model '{model_path}' {'(Mắt thần bật)' if use_web else ''}..."):
        try:
            resp_gen = requests.post(url_gen, headers=headers, json=payload, timeout=300) 
            if resp_gen.status_code == 200:
                try:
                    new_content = resp_gen.json()['candidates'][0]['content']['parts'][0]['text']
                    st.session_state.chat_history = [
                        {"role": "user", "parts": [{"text": prompt}]},
                        {"role": "model", "parts": [{"text": new_content}]}
                    ]
                    if che_do == "🎬 Kịch Bản Phim Tài Liệu (Chuẩn Silent Capital)":
                        st.session_state.sc_memory += f"\n\n{new_content}"
                        st.session_state.ai_result = st.session_state.sc_memory
                        st.session_state.sc_part += 1
                    else:
                        st.session_state.ai_result = new_content
                    st.rerun()
                except KeyError: st.error(f"❌ Nội dung nhạy cảm bị chặn. Chi tiết: {resp_gen.text}")
            else: st.error(f"❌ Lỗi API (Mã {resp_gen.status_code}): {resp_gen.text}")
        except Exception as e: st.error(f"Lỗi kết nối mạng: {e}")

# ================= GIAO DIỆN NHẬP LIỆU =================
tab1, tab2, tab3 = st.tabs(["📺 YouTube & Gợi Ý", "🌐 Phân Tích & So Sánh URL", "📁 Tệp Office & PDF"])

with tab1:
    che_do_yt = st.radio("Chọn chế độ phân tích:", ("🌍 Gợi ý tìm kiếm Đa quốc gia", "🔗 Bóc tách từ Link Kênh/Video (X-Ray Đối thủ)"))
    st.markdown("---")
    if che_do_yt == "🌍 Gợi ý tìm kiếm Đa quốc gia":
        col_opt1, col_opt2 = st.columns(2)
        with col_opt1: nguon_gợi_ý = st.selectbox("Chọn nền tảng:", ("YouTube", "Google Search"))
        with col_opt2: quoc_gia = st.selectbox("🌍 Chọn thị trường:", ("Việt Nam (VN)", "Mỹ (US)", "Anh (UK)", "Toàn cầu"))
        tu_khoa_nhap = st.text_input("🔑 Nhập từ khóa gốc:")
        market_map = {"Việt Nam (VN)": {"gl": "vn", "hl": "vi"}, "Mỹ (US)": {"gl": "us", "hl": "en"}, "Anh (UK)": {"gl": "gb", "hl": "en"}, "Toàn cầu": {"gl": "", "hl": "en"}}
        
        if st.button("🚀 Quét Từ Khóa"):
            if tu_khoa_nhap:
                gl, hl = market_map[quoc_gia]["gl"], market_map[quoc_gia]["hl"]
                danh_sach_truy_van = [tu_khoa_nhap] + [f"{tu_khoa_nhap} {chr(i)}" for i in range(97, 123)]
                tat_ca_suggests = []
                bar = st.progress(0)
                client_type = "yt" if nguon_gợi_ý == "YouTube" else "chrome"
                for i, q in enumerate(danh_sach_truy_van):
                    try:
                        url = f"https://suggestqueries.google.com/complete/search?client=firefox&ds={client_type}&q={q}&hl={hl}&gl={gl}" if nguon_gợi_ý == "YouTube" else f"https://suggestqueries.google.com/complete/search?client=firefox&q={q}&hl={hl}&gl={gl}"
                        res = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=5)
                        if res.status_code == 200: tat_ca_suggests.extend(res.json()[1])
                    except: pass
                    bar.progress((i + 1) / len(danh_sach_truy_van))
                    time.sleep(0.05)
                bar.empty()
                full_text = ' '.join(tat_ca_suggests)
                res_kw = trich_xuat_tu_khoa(lam_sach_text(full_text), selected_ngram)
                luu_ket_qua_vao_bo_nho(res_kw, full_text, f"{nguon_gợi_ý}_{market_map[quoc_gia]['gl']}")
            else: st.warning("Vui lòng nhập từ khóa!")
                
    else:
        link_yt = st.text_input("🔗 Nhập link Video YouTube:")
        quet_comment = st.checkbox("💬 Quét luôn cả bình luận", value=True)
        if st.button("🚀 Phân Tích & X-Ray Đối Thủ"):
            if not link_yt: st.warning("Vui lòng nhập link Video!")
            elif not yt_api_key: st.error("⚠️ Cần nhập YouTube API v3 Key!")
            else:
                with st.spinner("🕵️ Đang bóc tách dữ liệu..."):
                    try:
                        vid_id = re.search(r'(?:v=|\/)([0-9A-Za-z_-]{11}).*', link_yt).group(1)
                        full_content, xray = "", {}
                        vid_res = requests.get(f"https://www.googleapis.com/youtube/v3/videos?part=snippet,statistics&id={vid_id}&key={yt_api_key}").json()
                        if 'items' in vid_res and len(vid_res['items']) > 0:
                            v = vid_res['items'][0]
                            xray.update({'v_title': v['snippet']['title'], 'v_desc': v['snippet']['description'], 'v_tags': v['snippet'].get('tags', []), 'v_views': v['statistics'].get('viewCount', '0'), 'v_likes': v['statistics'].get('likeCount', '0'), 'v_comments': v['statistics'].get('commentCount', '0'), 'channel_id': v['snippet']['channelId'], 'channel_title': v['snippet']['channelTitle']})
                            thumbs = v['snippet']['thumbnails']
                            xray['v_thumb'] = thumbs.get('maxres', thumbs.get('high', thumbs.get('default')))['url']
                            full_content += f"{xray['v_title']} {xray['v_desc']} " + " ".join(xray['v_tags'])
                            ch_res = requests.get(f"https://www.googleapis.com/youtube/v3/channels?part=statistics&id={xray['channel_id']}&key={yt_api_key}").json()
                            if 'items' in ch_res:
                                xray.update({'c_subs': ch_res['items'][0]['statistics'].get('subscriberCount', 'Ẩn'), 'c_views': ch_res['items'][0]['statistics'].get('viewCount', '0'), 'c_videos': ch_res['items'][0]['statistics'].get('videoCount', '0')})
                        if quet_comment:
                            cmt_res = requests.get(f"https://www.googleapis.com/youtube/v3/commentThreads?part=snippet&videoId={vid_id}&maxResults=100&key={yt_api_key}").json()
                            if 'items' in cmt_res: full_content += " " + " ".join([i['snippet']['topLevelComment']['snippet']['textOriginal'] for i in cmt_res['items']])
                        xray['thumb_review'] = ai_cham_diem_thumbnail(xray['v_thumb'], xray['v_title'])
                        st.session_state.yt_xray_data = xray
                        luu_ket_qua_vao_bo_nho(trich_xuat_tu_khoa(lam_sach_text(full_content), selected_ngram), full_content, "youtube_xray_data")
                    except Exception as e: st.error(f"Lỗi: {e}")

with tab2:
    url1 = st.text_input("🔗 Nhập URL trang web:")
    if st.button("🚀 Phân Tích URL") and url1:
        res = requests.get(url1, headers={'User-Agent': 'Mozilla/5.0'}, timeout=10)
        soup = BeautifulSoup(res.text, 'html.parser')
        for script in soup(["script", "style"]): script.decompose()
        text1 = ' '.join([tag.get_text() for tag in soup.find_all(['p', 'h1', 'h2', 'article'])])
        luu_ket_qua_vao_bo_nho(trich_xuat_tu_khoa(lam_sach_text(text1), selected_ngram), text1, "Web_Keywords")

with tab3:
    # TÍNH NĂNG 3: BỂ CHỨA TÀI LIỆU PDF
    uploaded_file = st.file_uploader("Chọn file (pdf, xlsx, csv, docx, txt):", type=['pdf', 'xlsx', 'csv', 'docx', 'txt'])
    if st.button("🚀 Phân Tích Tệp") and uploaded_file:
        file_text = ""
        fname = uploaded_file.name.lower()
        if fname.endswith('.csv'): file_text = ' '.join(pd.read_csv(uploaded_file).astype(str).values.flatten())
        elif fname.endswith('.txt'): file_text = uploaded_file.read().decode('utf-8', errors='ignore')
        elif fname.endswith('.pdf'): 
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            file_text = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        luu_ket_qua_vao_bo_nho(trich_xuat_tu_khoa(lam_sach_text(file_text), selected_ngram), file_text, "document_file")

# ================= TỪ KHÓA & AI =================
if st.session_state.current_kw:
    st.markdown("---")
    df = pd.DataFrame(st.session_state.current_kw, columns=['Từ khóa', 'Tần suất'])
    st.dataframe(df.head(10), use_container_width=True)

    st.subheader("🤖 Xưởng Sản Xuất Nội Dung AI")
    
    # Nạp cả hành động mặc định và hành động tùy chỉnh
    default_actions = ["📝 Lập Dàn Ý SEO (Outline)", "✍️ Viết Bài Full Chuẩn SEO (1000+ từ)", "🎬 Kịch Bản Video Ngắn (TikTok/Reels/Shorts)", "🕵️ Tái tạo Kịch Bản YouTube Đối Thủ (Phân tích Link)", "🛒 Tối Ưu Gian Hàng TMĐT (Tiêu đề & Mô tả)", "🎬 Kịch Bản Phim Tài Liệu (Chuẩn Silent Capital)", "🧠 Phản Biện Kịch Bản (Script Doctor)", "📈 Dự Báo Thuật Toán YouTube (Pre-Publish)"]
    all_actions = default_actions + list(st.session_state.custom_prompts.keys())
    
    che_do_ai = st.selectbox("Chọn hành động:", all_actions)
    
    # Nút bật tắt mắt thần
    use_web_search = st.checkbox("🌐 Kích hoạt Mắt thần Internet (Quét Google thời gian thực để lấy trend mới nhất)")
    
    if st.button(f"✨ Kích Hoạt Trợ Lý AI"):
        xu_ly_ai_da_nang(che_do_ai, [i[0] for i in st.session_state.current_kw[:10]], st.session_state.current_text, False, use_web=use_web_search)
        
    if st.session_state.ai_result:
        st.success("✅ Trợ lý AI đã hoàn thành nhiệm vụ!")
        st.markdown(st.session_state.ai_result)
        
        # --- KHU VỰC TỔNG HỢP VÀ XUẤT FILE TỰ ĐỘNG ---
        st.markdown("---")
        st.subheader("📥 Tổng Hợp, Xuất File & Đọc Thử")
        
        full_export_text = "=== KẾT QUẢ TỔNG HỢP TỪ TRỢ LÝ AI ===\n\n"
        for msg in st.session_state.chat_history[1:]:
            role = "👤 Lệnh của bạn:" if msg["role"] == "user" else "🤖 Trợ lý AI:"
            full_export_text += f"{role}\n"
            for part in msg["parts"]:
                if "text" in part: full_export_text += part["text"] + "\n"
            full_export_text += "\n" + "-"*40 + "\n\n"

        if len(st.session_state.chat_history) > 0:
            col_ex1, col_ex2, col_ex3 = st.columns(3)
            with col_ex1:
                st.download_button(label="📄 Tải file Văn bản (.TXT)", data=full_export_text, file_name="Kich_Ban_Tong_Hop.txt", mime="text/plain")
            with col_ex2:
                doc_export = docx.Document()
                doc_export.add_heading('Kịch Bản Tổng Hợp AI', 0)
                doc_export.add_paragraph(full_export_text)
                buffer_docx = io.BytesIO()
                doc_export.save(buffer_docx)
                buffer_docx.seek(0)
                st.download_button(label="📘 Tải file Word (.DOCX)", data=buffer_docx, file_name="Kich_Ban_Tong_Hop.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # TÍNH NĂNG 1: TRỢ LÝ ĐỌC & THỬ GIỌNG (TTS)
            with col_ex3:
                if st.button("🔊 Nghe thử Kịch bản"):
                    with st.spinner("Đang tổng hợp giọng nói AI (có thể mất 1-2 phút tùy độ dài)..."):
                        try:
                            # Lấy đoạn phản hồi cuối cùng của AI để đọc
                            text_to_read = st.session_state.chat_history[-1]["parts"][0]["text"]
                            tts = gTTS(text=text_to_read, lang='vi')
                            audio_fp = io.BytesIO()
                            tts.write_to_fp(audio_fp)
                            st.audio(audio_fp, format='audio/mp3')
                        except Exception as e:
                            st.error(f"Lỗi tạo giọng nói: {e}")

        # --- KHU VỰC CHAT ĐÍNH KÈM TỆP (CÓ HỖ TRỢ PDF) ---
        st.markdown("---")
        st.subheader("💬 Trợ Lý Tinh Chỉnh & Phân Tích Kèm Tệp")
        
        for msg in st.session_state.chat_history[2:]:
            with st.chat_message("user" if msg["role"] == "user" else "assistant"):
                for part in msg["parts"]:
                    if "text" in part: st.markdown(part["text"])
                    if "inline_data" in part: st.image(base64.b64decode(part["inline_data"]["data"]), width=300)

        chat_file = st.file_uploader("📎 Đính kèm tệp (Ảnh, TXT, DOCX, CSV, PDF)", type=['png', 'jpg', 'jpeg', 'txt', 'docx', 'csv', 'pdf'], key="chat_upload")

        if prompt_chat := st.chat_input("Nhập yêu cầu tinh chỉnh hoặc phân tích tệp..."):
            user_parts = [{"text": prompt_chat}]
            
            if chat_file:
                fname = chat_file.name.lower()
                if fname.endswith(('.png', '.jpg', '.jpeg')):
                    img_b64 = base64.b64encode(chat_file.getvalue()).decode('utf-8')
                    mime = "image/png" if fname.endswith('.png') else "image/jpeg"
                    user_parts.append({"inline_data": {"mime_type": mime, "data": img_b64}})
                elif fname.endswith('.txt'): user_parts[0]["text"] += f"\n\n[DỮ LIỆU ĐÍNH KÈM]:\n{chat_file.getvalue().decode('utf-8', errors='ignore')}"
                elif fname.endswith('.csv'): user_parts[0]["text"] += f"\n\n[DỮ LIỆU ĐÍNH KÈM]:\n{' '.join(pd.read_csv(chat_file).astype(str).values.flatten())}"
                elif fname.endswith('.docx'): user_parts[0]["text"] += f"\n\n[DỮ LIỆU ĐÍNH KÈM]:\n{' '.join([p.text for p in docx.Document(chat_file).paragraphs])}"
                # Xử lý đọc PDF trong Chat
                elif fname.endswith('.pdf'): 
                    pdf_reader = PyPDF2.PdfReader(chat_file)
                    pdf_text = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
                    user_parts[0]["text"] += f"\n\n[DỮ LIỆU ĐÍNH KÈM TỪ PDF]:\n{pdf_text}"
            
            with st.chat_message("user"):
                st.markdown(prompt_chat)
                if chat_file:
                    if chat_file.name.lower().endswith(('.png', '.jpg', '.jpeg')): st.image(chat_file.getvalue(), width=300)
                    else: st.info(f"📄 Đã đính kèm tệp: {chat_file.name}")
            
            st.session_state.chat_history.append({"role": "user", "parts": user_parts})
            
            model_path = ai_model_choice.strip().replace("models/", "")
            url_chat = f"https://generativelanguage.googleapis.com/v1beta/models/{model_path}:generateContent"
            headers_chat = {'x-goog-api-key': gemini_api_key, 'Content-Type': 'application/json'}
            payload_chat = {"contents": st.session_state.chat_history}
            
            # Gắn mắt thần vào chat nếu đang bật
            if use_web_search: payload_chat["tools"] = [{"googleSearch": {}}]
            
            with st.chat_message("assistant"):
                with st.spinner("🤖 Đang phân tích..."):
                    try:
                        res_chat = requests.post(url_chat, headers=headers_chat, json=payload_chat, timeout=300)
                        if res_chat.status_code == 200:
                            reply = res_chat.json()['candidates'][0]['content']['parts'][0]['text']
                            st.markdown(reply)
                            st.session_state.chat_history.append({"role": "model", "parts": [{"text": reply}]})
                            st.rerun() 
                        else: st.error(f"Lỗi API: {res_chat.text}")
                    except Exception as e: st.error(f"Lỗi: {e}")
